"""
Generate a sample Parents Income Tax Return document for Aelin Galathynius.
Saves as a .docx file ready for upload testing.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os
from datetime import date

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Parents_Income_Tax_Return_Aelin_Galathynius.docx")


def set_cell_border(cell, **kwargs):
    """Set cell border properties."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = tc._tc.makeelement(qn('w:tcBorders'), {})
        tcPr.append(tcBorders)
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = tcBorders.makeelement(qn(f'w:{edge}'), edge_data)
            tcBorders.append(tag)


def make_cell(row, col, text, bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
    cell = row.cells[col]
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Arial"
    run.bold = bold
    if bg:
        shading = cell._tc.get_or_add_tcPr()
        shd = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): bg,
            qn('w:val'): 'clear',
        })
        shading.append(shd)
    return cell


def main():
    doc = Document()

    # ── Page setup ──
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ── Header ──
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run("REPUBLIC OF THE PHILIPPINES").bold = True
    header.runs[0].font.size = Pt(12)
    header.add_run("\nDEPARTMENT OF FINANCE").bold = True
    header.runs[1].font.size = Pt(11)
    header.add_run("\nBUREAU OF INTERNAL REVENUE").bold = True
    header.runs[2].font.size = Pt(11)

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ANNUAL INCOME TAX RETURN\nFor Individuals Earning Income Purely from Compensation\n(January 1, 2023 - December 31, 2023)")
    run.bold = True
    run.font.size = Pt(10)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("BIR Form No. 1700")
    run.bold = True
    run.font.size = Pt(11)

    # ── Taxpayer Info ──
    doc.add_paragraph()
    info_table = doc.add_table(rows=7, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    rows_data = [
        ("Last Name", "GALATHYNIUS", "First Name", "BRANDON"),
        ("Middle Name", "ASHIHARA", "Suffix", ""),
        ("Taxpayer Identification No. (TIN)", "314-592-687-000", "Date of Birth", "05/12/1978"),
        ("Registered Address", "7th House, Ayala, Zamboanga City, 7000", "Contact No.", "0917-131178"),
        ("Civil Status", "Married", "Spouse Name", "ELYNTHIA GALATHYNIUS"),
        ("Spouse TIN", "314-592-688-000", "Filing Status", "Married Filing Jointly"),
        ("Employer Name", "University of the Philippines", "Employer TIN", "000-916-715-000"),
    ]

    for i, (label1, val1, label2, val2) in enumerate(rows_data):
        make_cell(info_table.rows[i], 0, label1, bold=True, size=8, bg="F0F0F0")
        make_cell(info_table.rows[i], 1, val1, bold=False, size=9)
        make_cell(info_table.rows[i], 2, label2, bold=True, size=8, bg="F0F0F0")
        make_cell(info_table.rows[i], 3, val2, bold=False, size=9)

    # ── Income Section ──
    doc.add_paragraph()
    income_header = doc.add_paragraph()
    income_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = income_header.add_run("PART I — GROSS INCOME & ADJUSTMENTS")
    run.bold = True
    run.font.size = Pt(10)

    income_table = doc.add_table(rows=10, cols=3)
    income_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    income_data = [
        ("Item", "Description", "Amount (PHP)"),
        ("1", "Gross Compensation Income (Salary, Wages, etc.)", ""),
        ("  1a", "Basic Salary", "1,200,000.00"),
        ("  1b", "13th Month Pay & Other Benefits", "100,000.00"),
        ("  1c", "Allowances (Non-taxable up to limit)", "60,000.00"),
        ("2", "Less: Non-taxable / Exempt Compensation Income", ""),
        ("  2a", "13th Month Pay & Other Benefits Exclusion", "(90,000.00)"),
        ("  2b", "De Minimis Benefits", "(30,000.00)"),
        ("3", "Total Taxable Compensation Income", ""),
        ("  3a", "Net Taxable Income (1a + (1b - 2a) + (1c - 2b))", "1,240,000.00"),
    ]

    for i, (item, desc, amount) in enumerate(income_data):
        bold = item in ("1", "2", "3")
        make_cell(income_table.rows[i], 0, item, bold=bold, size=8, bg="F0F0F0" if bold else None)
        make_cell(income_table.rows[i], 1, desc, bold=bold, size=8)
        make_cell(income_table.rows[i], 2, amount, bold=bold, size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Tax Computation ──
    doc.add_paragraph()
    tax_header = doc.add_paragraph()
    run = tax_header.add_run("PART II — TAX COMPUTATION")
    run.bold = True
    run.font.size = Pt(10)

    tax_table = doc.add_table(rows=8, cols=2)
    tax_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tax_data = [
        ("DESCRIPTION", "AMOUNT (PHP)"),
        ("Net Taxable Income", "1,240,000.00"),
        ("Less: Personal Exemptions (Basic)", "(250,000.00)"),
        ("Net Taxable Income After Exemptions", "990,000.00"),
        ("Tax Due (Graduated Rates — BIR Table)", "162,500.00"),
        ("Less: Tax Credits / Withholding", ""),
        ("   Tax Withheld by Employer (Form 2316)", "(185,000.00)"),
        ("NET TAX PAYABLE (REFUNDABLE)", "(22,500.00)"),
    ]

    for i, (desc, amt) in enumerate(tax_data):
        bold = i == 0
        make_cell(tax_table.rows[i], 0, desc, bold=bold or i == len(tax_data)-1, size=8,
                  bg="F0F0F0" if i == 0 else None)
        make_cell(tax_table.rows[i], 1, amt, bold=bold or i == len(tax_data)-1, size=9,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ── Dependents Section ──
    doc.add_paragraph()
    dep_header = doc.add_paragraph()
    run = dep_header.add_run("PART III — DEPENDENT CHILDREN INFORMATION")
    run.bold = True
    run.font.size = Pt(10)

    dep_table = doc.add_table(rows=3, cols=5)
    dep_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    dep_data = [
        ("Name", "Date of Birth", "Relationship", "TIN", "Qualified Dependent"),
        ("AELIN ASHRYVER GALATHYNIUS", "01/16/2004", "Daughter", "N/A", "YES"),
        ("AIDAN ASHIHARA GALATHYNIUS", "03/22/2008", "Son", "N/A", "YES"),
    ]

    for i, row_data in enumerate(dep_data):
        for j, val in enumerate(row_data):
            bold = i == 0
            make_cell(dep_table.rows[i], j, val, bold=bold, size=8,
                      bg="F0F0F0" if i == 0 else None, align=WD_ALIGN_PARAGRAPH.CENTER if j >= 3 else WD_ALIGN_PARAGRAPH.LEFT)

    # ── Certification ──
    doc.add_paragraph()
    cert = doc.add_paragraph()
    cert.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cert.add_run("— TAXPAYER CERTIFICATION —").bold = True
    cert.runs[0].font.size = Pt(10)

    body = doc.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = body.add_run(
        "I declare under the penalties of perjury that this return has been made in good faith, "
        "verified by me, and to the best of my knowledge and belief, is true and correct, pursuant "
        "to the provisions of the National Internal Revenue Code, as amended, and the regulations "
        "issued under authority thereof."
    )
    run.font.size = Pt(9)

    # ── Signatures ──
    doc.add_paragraph()
    sig_table = doc.add_table(rows=4, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    sig_data = [
        ("", ""),
        ("_______________________________", "_______________________________"),
        ("BRANDON GALATHYNIUS", "ELYNTHIA GALATHYNIUS"),
        ("Taxpayer / Husband Signature", "Spouse Signature"),
    ]

    for i, (left, right) in enumerate(sig_data):
        make_cell(sig_table.rows[i], 0, left, bold=(i == 2), size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        make_cell(sig_table.rows[i], 1, right, bold=(i == 2), size=9,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Date Filed: April 15, 2024")
    run.font.size = Pt(9)

    # ── Attachments section ──
    doc.add_paragraph()
    attach = doc.add_paragraph()
    run = attach.add_run("ATTACHMENTS:")
    run.bold = True
    run.font.size = Pt(9)

    attachments = [
        "Certificate of Compensation Payment / Tax Withheld (BIR Form 2316)",
        "Waiver of Husband's Right to Claim Additional Exemption (if applicable)",
        "Approved Tax Debit Memo (if applicable)",
        "Proof of Foreign Tax Credits (if applicable)",
    ]
    for item in attachments:
        p = doc.add_paragraph(f"    □  {item}")
        p.paragraph_format.space_after = Pt(0)

    doc.save(OUTPUT_FILE)
    print(f"Created: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
